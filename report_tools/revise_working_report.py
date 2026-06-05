from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(
    r"C:\Users\user\Swinburne\Degree Year 3\SEM 2\COS30043 INTERFACE DESIGN AND DEVELOPMENT\IDD Custom Website.docx"
)
OUT_DIR = ROOT / "report_tools" / "final"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / "IDD_Custom_Website_Working_Report_Revised.docx"


def element_text(element):
    return "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()


def clean_heading_text(text):
    text = text.replace("\xa0", " ").strip()
    text = re.sub(r"^\d+(?:\.\d+)?\s+", "", text)
    return re.sub(r"\s+", " ", text).strip()


def normalise_key(text):
    return clean_heading_text(text).upper()


def is_heading_element(element, target):
    return normalise_key(element_text(element)) == target


def find_heading_index(children, target):
    for index, child in enumerate(children):
        if child.tag == qn("w:p") and is_heading_element(child, target):
            return index
    return None


def is_h1_element(element):
    if element.tag != qn("w:p"):
        return False
    p_style = element.find(".//w:pStyle", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
    if p_style is None:
        return False
    return p_style.get(qn("w:val")) == "Heading1"


def next_h1_index(children, start):
    for index in range(start + 1, len(children)):
        if is_h1_element(children[index]):
            return index
    return len(children)


def move_improvements_before_advanced(doc):
    body = doc.element.body
    children = list(body)
    advanced_start = find_heading_index(children, "ADVANCED FEATURES IMPLEMENTED")
    improvements_start = find_heading_index(children, "IMPROVEMENTS")
    if advanced_start is None or improvements_start is None:
        return False
    if improvements_start < advanced_start:
        return False

    improvements_end = next_h1_index(children, improvements_start)
    improvements_block = children[improvements_start:improvements_end]
    for child in improvements_block:
        body.remove(child)

    children = list(body)
    advanced_start = find_heading_index(children, "ADVANCED FEATURES IMPLEMENTED")
    for child in reversed(improvements_block):
        body.insert(advanced_start, child)
    return True


def set_paragraph_text(paragraph, text):
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")


def renumber_headings(doc):
    h1_numbers = {
        "INTRODUCTION": "1.0",
        "SYSTEM OVERVIEW": "2.0",
        "WIREFRAME DESIGN DIAGRAMS": "3.0",
        "WEB APPLICATION SCREENSHOTS": "4.0",
        "RESPONSIVE UI DESCRIPTION": "5.0",
        "USABILITY": "6.0",
        "ACCESSIBILITY EVALUATION": "7.0",
        "IMPROVEMENTS": "8.0",
        "ADVANCED FEATURES IMPLEMENTED": "9.0",
        "CONCLUSION": "10.0",
    }
    current_major = None
    subsection_counter = 0

    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        text = clean_heading_text(paragraph.text)
        key = text.upper()
        if style == "Heading 1" and key in h1_numbers:
            number = h1_numbers[key]
            current_major = number.split(".")[0]
            subsection_counter = 0
            set_paragraph_text(paragraph, f"{number} {key}")
        elif style == "Heading 2" and current_major is not None:
            subsection_counter += 1
            set_paragraph_text(paragraph, f"{current_major}.{subsection_counter} {text}")


def append_conclusion_and_references(doc):
    existing = {normalise_key(p.text) for p in doc.paragraphs if p.text.strip()}
    if "CONCLUSION" not in existing and "10.0 CONCLUSION" not in existing:
        doc.add_page_break()
        doc.add_heading("CONCLUSION", level=1)
        p = doc.add_paragraph(
            "PC Hardware Store demonstrates a full-stack, responsive Vue.js web application that supports both customer shopping workflows and administrator management workflows. The project includes a customer storefront, PC Builder, cart, wishlist, checkout, order history, profile management, reviews, and a separate administrator interface for products, orders, users, homepage products, profile management, and dashboard reporting."
        )
        p.paragraph_format.space_after = Pt(6)
        p = doc.add_paragraph(
            "The system applies interface design principles through consistent navigation, contextual grouping, row-column layouts, visual hierarchy, and role-based separation. It also demonstrates modern web development techniques including reusable Vue components, Pinia state management, Supabase database integration, RESTful CRUD operations, external API integration, live currency conversion, AI chat support, PC component recommendations, and responsive multi-device design."
        )
        p.paragraph_format.space_after = Pt(6)
        p = doc.add_paragraph(
            "Overall, the project satisfies the main requirements of the assignment while also identifying realistic future improvements such as stronger authentication, better field-level validation, deeper accessibility testing, pagination, and more formal user testing."
        )
        p.paragraph_format.space_after = Pt(6)

    existing = {normalise_key(p.text) for p in doc.paragraphs if p.text.strip()}
    if "REFERENCES" not in existing:
        doc.add_heading("REFERENCES", level=1)
        references = [
            "Vue.js. (2026). Introduction. https://vuejs.org/guide/introduction.html",
            "Vue Router. (2026). Vue Router Documentation. https://router.vuejs.org/",
            "Pinia. (2026). Pinia Documentation. https://pinia.vuejs.org/",
            "Vite. (2026). Vite Guide. https://vite.dev/guide/",
            "Supabase. (2026). Supabase Documentation. https://supabase.com/docs",
            "Google AI for Developers. (2026). Gemini API Documentation. https://ai.google.dev/gemini-api/docs",
            "Groq. (2026). Groq API Documentation. https://console.groq.com/docs",
            "Twelve Data. (2026). API Documentation. https://twelvedata.com/docs",
            "W3C Web Accessibility Initiative. (2026). Web Content Accessibility Guidelines. https://www.w3.org/WAI/standards-guidelines/wcag/",
            "MDN Web Docs. (2026). Responsive design. https://developer.mozilla.org/",
        ]
        list_style = "List Paragraph" if "List Paragraph" in [s.name for s in doc.styles] else None
        for reference in references:
            paragraph = doc.add_paragraph(reference, style=list_style)
            paragraph.paragraph_format.space_after = Pt(3)


def audit_doc(doc):
    full_text = "\n".join(p.text for p in doc.paragraphs)
    bad = [marker for marker in ("â", "Ã", "�") if marker in full_text]
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    return {
        "headings": len(headings),
        "tables": len(doc.tables),
        "bad_markers": bad,
        "first_headings": headings[:12],
        "last_headings": headings[-12:],
    }


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(str(SOURCE))
    append_conclusion_and_references(doc)
    moved = move_improvements_before_advanced(doc)
    renumber_headings(doc)
    doc.save(str(OUTPUT))

    out_doc = Document(str(OUTPUT))
    audit = audit_doc(out_doc)
    print(f"saved={OUTPUT}")
    print(f"moved_improvements_before_advanced={moved}")
    print(f"headings={audit['headings']}")
    print(f"tables={audit['tables']}")
    print(f"encoding_artifacts={audit['bad_markers'] or 'none'}")
    print("first_headings=")
    for heading in audit["first_headings"]:
        print(f"- {heading}")
    print("last_headings=")
    for heading in audit["last_headings"]:
        print(f"- {heading}")


if __name__ == "__main__":
    main()
