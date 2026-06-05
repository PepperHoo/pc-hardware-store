from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(r"C:\Users\user\Swinburne\Degree Year 3\SEM 2\COS30043 INTERFACE DESIGN AND DEVELOPMENT")
WORKSPACE = ROOT / "Custom Website"
DOCX_PATH = ROOT / "IDD Custom Website.docx"
PDF_PATH = ROOT / "Assignment" / "COS30043 S1 2026 Project.pdf"


def inspect_docx() -> str:
    document = Document(DOCX_PATH)
    lines = ["DOCX PARAGRAPHS"]
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            lines.append(f"[P{index:03d}] ({paragraph.style.name}) {text}")

    lines.append("\nDOCX TABLES")
    for table_index, table in enumerate(document.tables):
        lines.append(f"\n[TABLE {table_index}] {len(table.rows)} rows x {len(table.columns)} cols")
        for row_index, row in enumerate(table.rows):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            lines.append(f"  R{row_index:02d}: " + " | ".join(cells))
    return "\n".join(lines)


def inspect_pdf() -> str:
    reader = PdfReader(PDF_PATH)
    lines = [f"PDF PAGES: {len(reader.pages)}"]
    for page_index, page in enumerate(reader.pages, start=1):
        lines.append(f"\n--- PAGE {page_index} ---\n")
        lines.append(page.extract_text() or "")
    return "\n".join(lines)


def main() -> None:
    output_dir = WORKSPACE / "report_tools" / "extracted"
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "template.txt").write_text(inspect_docx(), encoding="utf-8")
    (output_dir / "brief.txt").write_text(inspect_pdf(), encoding="utf-8")
    print(output_dir)


if __name__ == "__main__":
    main()
