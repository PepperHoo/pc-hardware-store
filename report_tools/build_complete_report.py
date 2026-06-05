from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report_tools" / "final"
IMG_DIR = OUT_DIR / "wireframes"
SCREEN_DIR = ROOT / "report_tools" / "screenshots"
OUT_DOCX = OUT_DIR / "PC_Hardware_Project_Report_Complete.docx"

OUT_DIR.mkdir(parents=True, exist_ok=True)
IMG_DIR.mkdir(parents=True, exist_ok=True)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(84, 99, 117)
LIGHT_FILL = "F2F6FC"
HEADER_FILL = "E8EEF7"
BORDER = "C7D2E5"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths=None, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if widths:
                set_cell_width(cell, widths[cell_index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_font(run, size=9.5, color=NAVY)
            if header and row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, size=9, bold=True, color=NAVY)
            else:
                set_cell_shading(cell, "FFFFFF")


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, size=16, color=BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_font(run, size=13, color=BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        set_font(run, size=12, color=DARK_BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=11, color=NAVY, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_font(r, size=11, color=BLACK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    for run in p.runs:
        set_font(run, size=10.5, color=BLACK)
    if not p.runs:
        r = p.add_run(text)
        set_font(r, size=10.5, color=BLACK)
    else:
        p.runs[0].text = text
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9, color=MUTED, italic=True)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    style_table(table, widths=widths, header=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_border(cell, color="BFD3F2")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def maybe_add_picture(doc, image_path, width=6.1, caption=None):
    path = Path(image_path)
    if not path.exists():
        add_callout(doc, "Missing figure", f"Expected image file was not found: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        add_caption(doc, caption)


def load_font(size=28, bold=False):
    if ImageFont is None:
        return None
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_wireframe(path, title, elements, size=(1300, 760)):
    if Image is None:
        return
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    font_title = load_font(30, True)
    font_label = load_font(18, True)
    font_small = load_font(14, False)
    draw.rectangle([0, 0, size[0] - 1, size[1] - 1], outline=(150, 165, 185), width=3)
    draw.text((30, 24), title, fill=(15, 39, 70), font=font_title)
    for item in elements:
        x, y, w, h, label, fill = item
        draw.rounded_rectangle(
            [x, y, x + w, y + h],
            radius=16,
            fill=fill,
            outline=(90, 116, 150),
            width=2,
        )
        draw.text((x + 16, y + 14), label, fill=(15, 39, 70), font=font_label)
        if h > 80:
            draw.text((x + 16, y + 42), "responsive content block", fill=(82, 99, 120), font=font_small)
    img.save(path)


def generate_wireframes():
    draw_wireframe(
        IMG_DIR / "wireframe_customer_desktop.png",
        "Customer Home and Shopping Flow - Desktop Wireframe",
        [
            (40, 90, 1220, 70, "Navigation: menu, logo, links, search, currency, theme, wishlist, cart, profile", (242, 247, 255)),
            (40, 185, 1220, 210, "Hero section with product image, headline and primary actions", (232, 241, 255)),
            (40, 430, 280, 120, "Statistics", (247, 250, 255)),
            (350, 430, 280, 120, "Categories", (247, 250, 255)),
            (660, 430, 280, 120, "Hot selling products", (247, 250, 255)),
            (970, 430, 290, 120, "Latest products", (247, 250, 255)),
            (40, 585, 590, 120, "Footer shop links", (250, 252, 255)),
            (670, 585, 590, 120, "Account and contact links", (250, 252, 255)),
        ],
    )
    draw_wireframe(
        IMG_DIR / "wireframe_mobile.png",
        "Customer Mobile Wireframe",
        [
            (40, 90, 1220, 80, "Compact navigation with icons", (242, 247, 255)),
            (40, 190, 1220, 90, "Search row", (250, 252, 255)),
            (40, 305, 1220, 180, "Hero text and actions", (232, 241, 255)),
            (40, 510, 590, 150, "Statistic card", (247, 250, 255)),
            (670, 510, 590, 150, "Statistic card", (247, 250, 255)),
            (40, 690, 590, 50, "Product card column 1", (247, 250, 255)),
            (670, 690, 590, 50, "Product card column 2", (247, 250, 255)),
        ],
    )
    draw_wireframe(
        IMG_DIR / "wireframe_admin_desktop.png",
        "Administrator Dashboard - Desktop Wireframe",
        [
            (40, 90, 1220, 70, "Admin top navigation: logo, pages, currency, theme, profile", (242, 247, 255)),
            (40, 185, 1220, 100, "Page title and quick actions", (250, 252, 255)),
            (40, 315, 280, 110, "Products metric", (247, 250, 255)),
            (350, 315, 280, 110, "Orders metric", (247, 250, 255)),
            (660, 315, 280, 110, "Users metric", (247, 250, 255)),
            (970, 315, 290, 110, "Revenue metric", (247, 250, 255)),
            (40, 455, 760, 150, "Revenue chart", (232, 241, 255)),
            (830, 455, 430, 150, "Order status chart", (232, 241, 255)),
            (40, 635, 590, 90, "Latest products", (247, 250, 255)),
            (670, 635, 590, 90, "Latest orders", (247, 250, 255)),
        ],
    )
    draw_wireframe(
        IMG_DIR / "wireframe_pc_builder.png",
        "PC Builder Flow Wireframe",
        [
            (40, 90, 1220, 80, "Builder introduction and recommendation status", (242, 247, 255)),
            (40, 200, 800, 120, "Motherboard selection card", (247, 250, 255)),
            (40, 345, 800, 120, "Processor recommendation card", (247, 250, 255)),
            (40, 490, 800, 120, "RAM / GPU / Storage recommendation cards", (247, 250, 255)),
            (880, 200, 380, 410, "Your Build summary", (232, 241, 255)),
            (40, 635, 1220, 80, "Add build to cart summary", (250, 252, 255)),
        ],
    )


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = DARK_BLUE
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def add_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "PC Hardware Store - COS30043 Interface Design and Development"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "PC Hardware Project Report"
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("PC HARDWARE STORE")
    set_font(r, size=26, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Custom PC Component E-Commerce Website")
    set_font(r, size=16, color=BLUE, bold=True)

    add_callout(
        doc,
        "Project Focus",
        "A responsive Vue.js and Supabase web application for browsing PC components, building a compatible PC, placing orders, and managing store data through a separate administrator interface.",
    )

    rows = [
        ("Unit", "COS30043 Interface Design and Development"),
        ("Assessment", "Custom Website Project Report"),
        ("Student Name", "[Add your full name]"),
        ("Student ID", "[Add your student ID]"),
        ("Tutor / Lecturer", "[Add tutor or lecturer name]"),
        ("Hosted Website", "https://pc-hardware-store.vercel.app"),
        ("Video Demonstration", "[Add video demonstration link]"),
        ("Submission Date", "5 June 2026"),
    ]
    table = add_table(doc, ["Item", "Details"], rows, widths=[1.9, 4.5])
    table.rows[0].cells[0].paragraphs[0].runs[0].text = "Report Information"
    table.rows[0].cells[1].paragraphs[0].runs[0].text = ""

    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "Report Contents", 1)
    sections = [
        "1.0 Introduction",
        "2.0 System Overview",
        "3.0 Wireframe Design Diagrams",
        "4.0 Web Application Screenshots",
        "5.0 Responsive UI Description",
        "6.0 Usability Evaluation",
        "7.0 Accessibility Evaluation",
        "8.0 Proposed Improvements",
        "9.0 Advanced Features Implemented",
        "10.0 Conclusion",
        "References",
        "Appendix A: Vue Implementation Evidence",
        "Appendix B: Usability and Accessibility Test Checklist",
    ]
    for section in sections:
        add_bullet(doc, section)
    doc.add_page_break()


def add_intro(doc):
    add_heading(doc, "1.0 Introduction", 1)
    add_para(
        doc,
        "PC Hardware Store is a full-stack e-commerce web application designed for customers who want to purchase individual PC components or assemble a complete personal computer. PC component shopping can be difficult because processors, motherboards, RAM, graphics cards, storage, power supplies, coolers, cases, and RGB accessories must be selected carefully to create a functioning build. The website addresses this problem by combining a conventional online store with a guided PC Builder that supports component selection from the actual store catalogue.",
    )
    add_para(
        doc,
        "The application was developed using Vue.js 3 with Vite, Vue Router, Pinia state management, Supabase database storage, Supabase Storage, and serverless API endpoints. It contains a customer-facing storefront and a separate administrator interface. Customers can browse products, view details, compare products, add items to a wishlist, manage a cart, place orders, print receipts, review products, and update profile information. Administrators can manage products, stock, orders, users, homepage product lists, profile details, and dashboard analytics.",
    )
    add_para(
        doc,
        "The system integrates external services to increase the quality of the user experience. A Google Gemini powered chat assistant supports customer questions about PC components, while a Groq powered recommendation endpoint supports the PC Builder by suggesting remaining component categories after the user selects a first item. A Twelve Data endpoint is used for live foreign exchange rates so product and order prices can be displayed in MYR, JPY, KRW, USD, SGD, and EUR when the required API key and quota are available.",
    )
    add_para(
        doc,
        "A major design objective was to make the website responsive and usable across desktop, tablet, and mobile devices. The implementation uses flexible grids, responsive navigation, adaptive product cards, media queries, light and dark themes, visible feedback, and touch-friendly controls. This report explains the system structure, design decisions, responsive layout techniques, usability evaluation, accessibility evaluation, proposed improvements, and advanced features implemented in the project.",
    )


def add_system_overview(doc):
    add_heading(doc, "2.0 System Overview", 1)
    add_heading(doc, "2.1 System Purpose", 2)
    add_para(
        doc,
        "The purpose of PC Hardware Store is to support the complete customer journey for purchasing PC components. The site is not only a product catalogue; it also assists customers who may not fully understand component compatibility. By dividing PC building into clear categories and offering AI-assisted recommendations, the system reduces the difficulty of selecting parts for a working computer.",
    )
    add_table(
        doc,
        ["Role", "Main Goals", "Key Pages"],
        [
            ("Customer", "Browse components, compare products, build a PC, manage wishlist/cart, checkout, review orders.", "Home, Products, Product Details, PC Builder, Wishlist, Cart, Checkout, Orders, Profile"),
            ("Administrator", "Manage inventory, stock, users, orders, homepage listings, profile, and performance information.", "Dashboard, Products, Orders, Order Details, Users, Homepage Editor, Admin Profile"),
        ],
        widths=[1.2, 3.0, 2.2],
    )

    add_heading(doc, "2.2 Architecture", 2)
    add_para(
        doc,
        "The application follows a client-server architecture. The frontend is a Vue.js single-page application built with Vite. Vue Router manages navigation and route protection, while Pinia stores manage shared state such as cart items, wishlist items, comparison data, and currency selection. Supabase stores persistent data and product/profile images. Vercel serverless functions connect the frontend to external AI and currency services without exposing private API keys to the browser.",
    )
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Frontend", "Vue.js 3, Vite, Vue Router", "Builds the interactive single-page interface and page navigation."),
            ("State", "Pinia", "Shares cart, wishlist, comparison, and currency state across components."),
            ("Backend Data", "Supabase REST API", "Stores and retrieves products, users, orders, reviews, cart items, wishlist items, and homepage product selections."),
            ("File Storage", "Supabase Storage", "Stores product images and profile images."),
            ("Serverless APIs", "Vercel API routes", "Protects API keys and handles Google AI, Groq recommendations, and Twelve Data exchange rates."),
        ],
        widths=[1.25, 1.65, 3.5],
    )

    add_heading(doc, "2.3 Customer-Facing Features", 2)
    customer_features = [
        "Home page with brand introduction, product categories, store statistics, hot-selling products, latest products, and direct calls to action.",
        "Products page with category filtering, product search, sorting, grid/list display, quantity controls, details, wishlist, comparison, and add-to-cart actions.",
        "Product detail page with image, category, stock, price, description, quantity controls, related products, wishlist action, comparison action, and customer reviews.",
        "PC Builder page with component categories and smart recommendations for remaining parts after a first component is selected.",
        "Cart and checkout flow with quantity updates, shipping method, payment method, order summary, stock validation, and automatic stock deduction after a successful order.",
        "Wishlist, order history, receipt printing, profile statistics, total spending, recent orders, and profile editing.",
    ]
    for item in customer_features:
        add_bullet(doc, item)

    add_heading(doc, "2.4 Administration Features", 2)
    admin_features = [
        "Admin Dashboard displays total products, total orders, total users, monthly revenue, total revenue, revenue chart, order status chart, products by category, latest products, and latest orders.",
        "Admin Products supports adding, editing, deleting, uploading product images, filtering products, and monitoring low stock.",
        "Admin Orders allows the administrator to view orders, update order status, open order details, and inspect purchased items and delivery information.",
        "Admin Users summarises total users, administrators, and customers, and provides search for user management.",
        "Admin Homepage Editor manages the hot-selling and latest product sections shown on the customer homepage.",
        "Admin Profile supports profile editing, profile image change, and logout.",
    ]
    for item in admin_features:
        add_bullet(doc, item)

    add_heading(doc, "2.5 Database and RESTful API Integration", 2)
    add_para(
        doc,
        "Supabase provides persistent data storage for the project. The frontend communicates with Supabase using asynchronous RESTful API requests through a reusable API helper module. The main tables include products, users, orders, reviews, cart_items, wishlist_items, and homepage selections. This makes core user data available across page refreshes, sign-out, and login from another device.",
    )
    add_table(
        doc,
        ["Data Area", "Stored Information", "CRUD Example"],
        [
            ("Products", "Name, price, category, stock, description, image.", "Admin creates/updates/deletes products; checkout updates stock."),
            ("Orders", "Customer, items, totals, shipping, payment, status.", "Checkout creates orders; admin updates order status."),
            ("Cart", "User email, product ID, product snapshot, quantity.", "Customer adds, updates, removes, and clears cart items."),
            ("Wishlist", "User email, product ID, product snapshot.", "Customer saves or removes wishlist products."),
            ("Reviews", "Product ID, user, rating, comment.", "Customer submits product review data."),
        ],
        widths=[1.25, 2.4, 2.75],
    )

    add_heading(doc, "2.6 External Service Integration", 2)
    add_table(
        doc,
        ["Service", "Used For", "Implementation Note"],
        [
            ("Google Gemini", "Customer AI chat assistant.", "Called through /api/chat so the key is kept in server environment variables."),
            ("Groq", "PC Builder smart component recommendations.", "Called through /api/pc-recommendations; returned IDs are validated against the current product catalogue."),
            ("Twelve Data", "Live foreign exchange rates.", "Called through /api/exchange-rates; rates are cached and reused if the provider is temporarily unavailable."),
        ],
        widths=[1.25, 2.1, 3.05],
    )
    add_callout(
        doc,
        "Implementation boundary",
        "The AI recommendation feature provides assisted suggestions from the store catalogue. It should not be described as a guaranteed engineering compatibility checker because final hardware compatibility still depends on manufacturer specifications such as CPU socket, RAM generation, PSU wattage, GPU length, and case clearance.",
    )

    add_heading(doc, "2.7 Interface Design Approach", 2)
    add_para(
        doc,
        "The interface applies core design principles required by the project brief. Contextual grouping is used to keep related actions close together, such as product quantity controls near add-to-cart buttons and order status controls near order actions. A row-column grid is used for product cards, dashboard statistics, forms, charts, and tables. Visual hierarchy is created through large page headings, smaller section headings, highlighted prices, status labels, and primary action buttons. Navigation is consistent within each role so that users can move through the customer side or administrator side without confusion.",
    )

    add_heading(doc, "2.8 Vue Framework Techniques", 2)
    add_para(
        doc,
        "The implementation uses Vue features throughout the project. Reusable components such as Navbar, AdminNavbar, Footer, Toast, AiChatWidget, CompareBar, ProductReviews, and loading skeletons reduce duplicated code. Directives such as v-if and v-for control conditional rendering and repeated product/status lists. v-model is used for form inputs and search fields. Dynamic bindings such as :class and :disabled adjust styling and button availability based on state, stock, selected currency, theme, and user role. The Composition API, computed values, watchers, and asynchronous fetch calls are used to retrieve data and update the interface reactively.",
    )


def add_wireframes(doc):
    add_heading(doc, "3.0 Wireframe Design Diagrams", 1)
    add_para(
        doc,
        "The wireframes below show the intended structure before final styling. They focus on layout, grouping, navigation, and information hierarchy rather than colour or final visual details.",
    )
    maybe_add_picture(
        doc,
        IMG_DIR / "wireframe_customer_desktop.png",
        width=6.2,
        caption="Figure 3.1 Customer desktop wireframe showing the main shopping flow.",
    )
    maybe_add_picture(
        doc,
        IMG_DIR / "wireframe_mobile.png",
        width=6.2,
        caption="Figure 3.2 Customer mobile wireframe showing compact navigation and stacked content.",
    )
    maybe_add_picture(
        doc,
        IMG_DIR / "wireframe_admin_desktop.png",
        width=6.2,
        caption="Figure 3.3 Administrator dashboard wireframe showing metrics, charts, and quick actions.",
    )
    maybe_add_picture(
        doc,
        IMG_DIR / "wireframe_pc_builder.png",
        width=6.2,
        caption="Figure 3.4 PC Builder wireframe showing component selection and build summary.",
    )


def add_screenshots(doc):
    add_heading(doc, "4.0 Web Application Screenshots", 1)
    add_para(
        doc,
        "This section presents screenshots from the completed application. The screenshots demonstrate the visual design, customer workflow, administrator workflow, and responsive implementation.",
    )
    shots = [
        ("home-desktop.png", "Figure 4.1 Customer homepage desktop view with hero content, navigation, currency control, theme toggle, wishlist, cart, profile, and AI assistant."),
        ("products-desktop.png", "Figure 4.2 Products page desktop view with product browsing, filtering, search, quantity control, and add-to-cart actions."),
        ("pc-builder-desktop.png", "Figure 4.3 PC Builder desktop view with component categories, recommendations, and build summary."),
        ("admin-dashboard-desktop.png", "Figure 4.4 Admin dashboard desktop view with store metrics, revenue information, charts, and latest data."),
        ("admin-products-desktop.png", "Figure 4.5 Admin Products view showing inventory management, add-product form, product cards, stock, edit, and delete actions."),
    ]
    for file_name, caption in shots:
        maybe_add_picture(doc, SCREEN_DIR / file_name, width=6.2, caption=caption)


def add_responsive(doc):
    add_heading(doc, "5.0 Responsive UI Description", 1)
    add_heading(doc, "5.1 Responsive Design Approach", 2)
    add_para(
        doc,
        "The application was designed to remain usable on desktop computers, tablets, and mobile phones. Responsive design is essential for this project because users may browse products, compare parts, build a PC, manage a cart, or check an order on different devices. The website uses CSS Grid, Flexbox, relative widths, minmax(), auto-fit, auto-fill, clamp(), media queries, and controlled image containers to reorganise content based on the available screen width.",
    )
    add_table(
        doc,
        ["Device Type", "Example Viewport", "Responsive Behaviour"],
        [
            ("Desktop", "1440 x 1000", "Full navigation, wide search bar, sidebars, multi-column cards, and dashboard charts."),
            ("Tablet", "820 x 1180", "Reduced columns, simplified spacing, stacked sections where needed, and compact navigation."),
            ("Mobile", "390 x 844", "Icon-first navigation, search on its own row, stacked sections, touch-friendly controls, and simplified cards."),
        ],
        widths=[1.3, 1.35, 3.75],
    )

    add_heading(doc, "5.2 Customer Homepage Responsiveness", 2)
    add_para(
        doc,
        "The homepage uses a large hero image and clear headline on desktop. On mobile, the navigation changes to icons, the search bar moves below the main row, and the hero content becomes centred and compact. Store statistics and product sections stack vertically so text and buttons remain readable.",
    )
    maybe_add_picture(doc, SCREEN_DIR / "home-desktop.png", width=6.1, caption="Figure 5.1 Homepage desktop layout.")
    maybe_add_picture(doc, SCREEN_DIR / "home-tablet.png", width=4.2, caption="Figure 5.2 Homepage tablet layout.")
    maybe_add_picture(doc, SCREEN_DIR / "home-mobile.png", width=2.3, caption="Figure 5.3 Homepage mobile layout.")

    add_heading(doc, "5.3 Product Listing Responsiveness", 2)
    add_para(
        doc,
        "The Products page uses a sidebar for categories and sorting on large screens. On tablet and mobile sizes, the sidebar can be hidden so that product cards receive more width. The product grid adjusts the number of columns to prevent overlap and keeps product images, prices, quantity controls, and add-to-cart buttons aligned.",
    )
    maybe_add_picture(doc, SCREEN_DIR / "products-desktop.png", width=6.1, caption="Figure 5.4 Products page desktop layout.")
    maybe_add_picture(doc, SCREEN_DIR / "products-mobile.png", width=2.6, caption="Figure 5.5 Products page mobile layout.")

    add_heading(doc, "5.4 PC Builder Responsiveness", 2)
    add_para(
        doc,
        "The PC Builder uses a two-column layout on desktop, with component selection on the left and the build summary on the right. On mobile, the builder changes to a single-column flow. This prevents the summary card from squeezing the recommendation cards and allows each component card to remain touch-friendly.",
    )
    maybe_add_picture(doc, SCREEN_DIR / "pc-builder-desktop.png", width=6.1, caption="Figure 5.6 PC Builder desktop layout.")
    maybe_add_picture(doc, SCREEN_DIR / "pc-builder-mobile.png", width=2.4, caption="Figure 5.7 PC Builder mobile layout.")

    add_heading(doc, "5.5 Administrator Responsiveness", 2)
    add_para(
        doc,
        "The administrator interface uses a top navigation bar. On desktop, labels and management controls are visible. On smaller screens, the navigation becomes icon-based and content sections stack. Product management cards use responsive grids, while order and user tables are placed inside scrollable containers to preserve readable column structure.",
    )
    maybe_add_picture(doc, SCREEN_DIR / "admin-dashboard-desktop.png", width=6.1, caption="Figure 5.8 Admin dashboard desktop layout.")
    maybe_add_picture(doc, SCREEN_DIR / "admin-dashboard-mobile.png", width=2.4, caption="Figure 5.9 Admin dashboard mobile layout.")

    add_heading(doc, "5.6 Responsive Testing Summary", 2)
    add_table(
        doc,
        ["Test Area", "Expected Result", "Observed Result"],
        [
            ("Navigation", "Important controls remain accessible.", "Customer and admin navigation switch to compact/icon layouts on mobile."),
            ("Product Cards", "Images, names, prices, and buttons remain aligned.", "Cards resize and reorganise without covering text."),
            ("Forms", "Inputs remain usable on small screens.", "Checkout, profile, and admin forms stack into one column where needed."),
            ("Tables", "Complex data remains readable.", "Admin order and user tables use responsive containers."),
            ("Themes", "Light and dark mode remain readable.", "Theme variables update backgrounds, text, cards, dropdowns, and controls."),
        ],
        widths=[1.4, 2.3, 2.7],
    )


def add_usability(doc):
    add_heading(doc, "6.0 Usability Evaluation", 1)
    add_para(
        doc,
        "The usability evaluation used a task-based walkthrough and common usability principles, including consistency, visibility of system status, error prevention, recognition rather than recall, user control, and efficiency of use. The evaluation considered both customer and administrator tasks.",
    )
    add_heading(doc, "6.1 Navigation Efficiency", 2)
    add_para(
        doc,
        "The customer navigation provides direct access to Products, PC Builder, search, currency, theme, wishlist, cart, and profile. The administrator navigation provides direct access to Dashboard, Products, Orders, Users, Homepage, and Admin Profile. This separation makes each role easier to understand and prevents customers and administrators from being mixed into the wrong workflow.",
    )
    add_heading(doc, "6.2 Product Browsing and Selection", 2)
    add_para(
        doc,
        "Customers can search, filter, sort, compare, save to wishlist, view details, adjust quantity, and add to cart. Stock information is visible before purchase, and unavailable products can have disabled purchase actions. This improves error prevention because users receive feedback before they attempt an impossible purchase.",
    )
    add_heading(doc, "6.3 PC Builder Usability", 2)
    add_para(
        doc,
        "The PC Builder divides the complex task of assembling a computer into recognisable categories. The Your Build summary keeps selected components visible, reducing memory load. After a first selection, the recommendation feature suggests remaining parts from the store catalogue. This supports less experienced users, although final compatibility should still be confirmed using manufacturer specifications.",
    )
    add_heading(doc, "6.4 Checkout and Form Usability", 2)
    add_para(
        doc,
        "The checkout page is grouped into delivery information, shipping method, payment method, and order summary. Product management and profile forms use repeated input styling and clear primary actions. Feedback is provided through toast messages and disabled states. A future improvement is to show validation messages directly beside the affected form field.",
    )
    add_heading(doc, "6.5 Administrator Usability", 2)
    add_para(
        doc,
        "The dashboard gives administrators a quick overview of inventory, users, orders, and revenue. Product management cards and order status controls are placed close to the relevant data, supporting efficient repeated management tasks. However, larger catalogues would benefit from pagination, advanced filtering, and bulk actions.",
    )
    add_table(
        doc,
        ["Usability Issue", "User Impact", "Suggested Improvement"],
        [
            ("Long product lists", "Users may take longer to locate products.", "Add pagination, price/spec filters, and visible active filters."),
            ("Mobile PC Builder scrolling", "Completing a full build takes longer on small screens.", "Add progress indicators and collapsible completed sections."),
            ("General form errors", "Users may not know which field caused the issue.", "Display field-level validation and focus the first invalid input."),
            ("Technical API messages", "Non-technical users may not understand recovery steps.", "Translate errors into plain language and provide retry/manual alternatives."),
            ("Destructive admin actions", "Products or users could be removed accidentally.", "Add confirmation dialogs for delete actions."),
        ],
        widths=[1.55, 2.2, 2.65],
    )


def add_accessibility(doc):
    add_heading(doc, "7.0 Accessibility Evaluation", 1)
    add_para(
        doc,
        "The accessibility evaluation was based on WCAG principles, focusing on whether the interface is perceivable, operable, understandable, and robust for users with different visual, physical, cognitive, and technical needs. The evaluation does not claim full WCAG compliance, but identifies current support and future improvements.",
    )
    add_heading(doc, "7.1 Current Accessibility Support", 2)
    features = [
        "Standard HTML buttons, links, input fields, and form controls are used for most interactions, supporting normal keyboard navigation.",
        "Icon-based controls such as menu, wishlist, cart, profile, AI chat open/close, and send message include accessible labels in key components.",
        "Toast feedback uses live status behaviour so feedback can be announced without forcing the user away from the current task.",
        "Product images use product-name alternative text, while decorative imagery can be kept from adding unnecessary screen reader noise.",
        "Status colours are supported by text labels such as Pending, Delivered, Rejected, In Stock, and Out of Stock.",
        "Responsive layouts and touch-friendly controls support mobile users and users who zoom the browser.",
    ]
    for item in features:
        add_bullet(doc, item)

    add_heading(doc, "7.2 Accessibility for Different Users", 2)
    add_table(
        doc,
        ["User Need", "Current Support", "Remaining Risk"],
        [
            ("Low vision", "Light/dark theme, clear headings, visible buttons, responsive layout.", "Some secondary text and disabled text should be contrast-tested."),
            ("Colour vision deficiency", "Status colours are paired with text labels.", "Avoid relying on colour alone in charts and status badges."),
            ("Keyboard-only use", "Standard controls and search inputs are keyboard accessible.", "Custom dropdowns, modals, and comparison controls need full keyboard testing."),
            ("Screen reader use", "Alt text, aria labels, and live status feedback are included.", "Some icons/category symbols may need clearer accessible names."),
            ("Mobile/touch use", "Compact navigation, larger buttons, and stacked layouts.", "Admin tables may still require horizontal scrolling."),
            ("Cognitive support", "Contextual grouping and repeated design patterns.", "Complex PC recommendations should include clearer explanation notes."),
        ],
        widths=[1.35, 2.7, 2.35],
    )

    add_heading(doc, "7.3 Accessibility Issues Identified", 2)
    add_table(
        doc,
        ["Issue", "Potential Impact", "Improvement"],
        [
            ("Low-contrast secondary text", "Users with low vision may miss supporting information.", "Use a contrast checker and adjust colour tokens."),
            ("Custom controls may lack full keyboard support", "Keyboard-only users may struggle with dropdowns or dialogs.", "Support Tab, Enter, Space, Escape, and arrow keys where appropriate."),
            ("General error notifications", "Screen reader users may not know which field needs correction.", "Connect error messages to inputs using aria-describedby."),
            ("Motion effects", "Motion-sensitive users may feel discomfort.", "Respect prefers-reduced-motion and reduce non-essential animations."),
            ("Technical AI/API errors", "Users may not understand what to do next.", "Show simple messages with retry and manual selection alternatives."),
        ],
        widths=[1.75, 2.3, 2.35],
    )


def add_improvements(doc):
    add_heading(doc, "8.0 Proposed Improvements", 1)
    add_para(
        doc,
        "The following improvements are proposed based on the usability and accessibility evaluation. They are prioritised by impact on task success, inclusiveness, security, and maintainability.",
    )
    add_table(
        doc,
        ["Priority", "Improvement", "Expected Benefit"],
        [
            ("High", "Add field-level validation, clear error messages, and focus movement to the first invalid field.", "Improves form success and accessibility."),
            ("High", "Complete keyboard testing for custom dropdowns, modals, comparison tools, and mobile navigation.", "Supports keyboard-only users."),
            ("High", "Audit colour contrast across light and dark themes.", "Improves readability for low-vision users."),
            ("High", "Migrate custom user authentication to Supabase Auth and strengthen Row Level Security.", "Improves data protection and account security."),
            ("Medium", "Add pagination, advanced filters, active filter indicators, and reset controls.", "Improves product discovery and admin efficiency."),
            ("Medium", "Add delete confirmation dialogs for products, users, and homepage items.", "Reduces accidental data loss."),
            ("Medium", "Add PC Builder explanation notes for socket, chipset, RAM generation, PSU wattage, and case clearance.", "Helps users understand why parts are recommended."),
            ("Medium", "Add reduced-motion support.", "Improves comfort for motion-sensitive users."),
            ("Low", "Optimise images, lazy load large media, debounce search, and split large route components.", "Improves loading speed and responsiveness."),
            ("Low", "Conduct formal user testing with users of different PC knowledge levels.", "Provides evidence for future design decisions."),
        ],
        widths=[0.75, 3.1, 2.55],
    )
    add_callout(
        doc,
        "Most important future improvement",
        "The most important improvement is replacing the custom users table with Supabase Auth and stricter Row Level Security. This would make the application more suitable for production use because users would only be able to access their own cart, wishlist, profile, and order data.",
    )


def add_advanced(doc):
    add_heading(doc, "9.0 Advanced Features Implemented", 1)
    add_para(
        doc,
        "The project includes several advanced features that extend it beyond a simple static website. These features demonstrate external API integration, database persistence, reusable component design, state management, responsive design, role-based routing, and dynamic user feedback.",
    )
    add_table(
        doc,
        ["Advanced Feature", "Implementation", "Value Added"],
        [
            ("Supabase CRUD and Storage", "Products, users, orders, reviews, cart items, wishlist items, and images are stored through Supabase REST and Storage.", "Provides persistent data and functional admin/customer workflows."),
            ("Google AI Chat Assistant", "Customer chat requests are sent through /api/chat to Google Gemini when configured.", "Helps customers ask PC component and build-planning questions."),
            ("Groq PC Recommendations", "The PC Builder sends selected parts and catalogue data to /api/pc-recommendations.", "Suggests remaining components from the real store catalogue."),
            ("Live Currency Conversion", "Currency store calls /api/exchange-rates and formats values across pages.", "Allows customer and admin prices to change between MYR, JPY, KRW, USD, SGD, and EUR."),
            ("Pinia State Management", "Cart, wishlist, comparison, and currency state are shared across components.", "Keeps navigation counts, totals, and page data consistent."),
            ("Role-Based Routing", "Vue Router guards use route meta fields such as requiresAuth and requiresAdmin.", "Separates customer and administrator workflows."),
            ("Stock Validation and Deduction", "Checkout checks current stock and updates product quantity after purchase.", "Reduces overselling and keeps inventory more accurate."),
            ("Comparison, Wishlist, Reviews", "Customers can compare products, save wishlist items, and submit reviews.", "Supports informed purchase decisions."),
            ("Admin Analytics", "Dashboard charts and cards summarise products, orders, users, revenue, latest products, and latest orders.", "Improves administrator monitoring."),
            ("Dynamic Themes and Responsive UI", "Theme variables and responsive CSS adapt colours and layouts across devices.", "Improves visual comfort and multi-device usability."),
        ],
        widths=[1.65, 2.65, 2.1],
    )
    add_heading(doc, "9.1 AI Recommendation Workflow", 2)
    add_para(
        doc,
        "When a user selects the first component in the PC Builder, the frontend sends the selected component, available categories, and current product catalogue to a serverless endpoint. The server calls the configured AI provider and asks for category-based recommendations. The response is then validated before being displayed. Validation is important because only product IDs that exist in the Supabase catalogue should appear in the user interface.",
    )
    add_heading(doc, "9.2 Currency Workflow", 2)
    add_para(
        doc,
        "Product prices are stored in MYR as the base currency. The currency Pinia store retrieves exchange rates from the serverless exchange-rate endpoint. When the selected currency changes, the same format function is used across product listings, product details, cart, checkout, profile, order history, admin products, admin orders, and dashboard revenue values. This makes currency display consistent across the whole system.",
    )
    add_heading(doc, "9.3 Component Reuse", 2)
    add_para(
        doc,
        "Reusable components reduce duplicated interface code. The same navigation, footer, toast, AI chat, comparison, review, and loading components can be maintained centrally. This is important for a larger website because design updates, accessibility improvements, and bug fixes can be applied to shared components instead of being repeated manually on every page.",
    )


def add_conclusion(doc):
    add_heading(doc, "10.0 Conclusion", 1)
    add_para(
        doc,
        "PC Hardware Store meets the project goal of creating a full-stack, responsive, interactive web application using contemporary web technologies. The system provides more than eight interconnected pages, a customer storefront, a role-separated administrator interface, Supabase-backed CRUD functionality, persistent cart and wishlist storage, product and order management, responsive layouts, theme switching, live currency conversion, AI chat support, and AI-assisted PC component recommendations.",
    )
    add_para(
        doc,
        "The interface applies key design principles such as contextual grouping, consistent navigation, visual hierarchy, row-column grids, and repeated card/form patterns. The responsive implementation allows the same application to work across desktop, tablet, and mobile screens. The usability and accessibility evaluation also identifies realistic limitations, including the need for stronger validation, further keyboard testing, improved contrast auditing, more secure authentication, pagination, and formal user testing.",
    )
    add_para(
        doc,
        "Overall, the project demonstrates a functional and polished e-commerce experience for PC hardware components while also showing independent exploration through database integration, external APIs, state management, administrator analytics, responsive design, and advanced user support features.",
    )


def add_references(doc):
    add_heading(doc, "References", 1)
    refs = [
        "Vue.js. (2026). Introduction. https://vuejs.org/guide/introduction.html",
        "Vue Router. (2026). Vue Router Documentation. https://router.vuejs.org/",
        "Pinia. (2026). Pinia Documentation. https://pinia.vuejs.org/",
        "Supabase. (2026). Supabase Documentation. https://supabase.com/docs",
        "Google AI for Developers. (2026). Gemini API Documentation. https://ai.google.dev/gemini-api/docs",
        "Groq. (2026). Groq API Documentation. https://console.groq.com/docs",
        "Twelve Data. (2026). API Documentation. https://twelvedata.com/docs",
        "W3C Web Accessibility Initiative. (2026). Web Content Accessibility Guidelines (WCAG). https://www.w3.org/WAI/standards-guidelines/wcag/",
        "MDN Web Docs. (2026). Responsive Design. https://developer.mozilla.org/",
    ]
    for ref in refs:
        add_bullet(doc, ref)


def add_appendices(doc):
    add_heading(doc, "Appendix A: Vue Implementation Evidence", 1)
    add_table(
        doc,
        ["Vue Technique", "Example in Project", "Purpose"],
        [
            ("v-for", "Product lists, category lists, order rows, currency options, chat messages.", "Renders repeated data from arrays."),
            ("v-if", "Search dropdowns, loading states, empty states, AI chat panel, selected build sections.", "Shows or hides interface elements based on state."),
            ("v-model", "Search fields, profile forms, admin product forms, review text, AI chat input.", "Creates two-way binding for form values."),
            (":class", "Active navigation links, selected currency, stock status, theme states.", "Applies dynamic styling based on data."),
            (":disabled", "Out-of-stock buttons, invalid submit buttons, unavailable comparison actions.", "Prevents actions that cannot be completed."),
            ("computed", "Cart count, totals, filtered products, selected route/profile path.", "Keeps derived values reactive and consistent."),
            ("watch", "Currency changes, theme changes, route/data refreshes.", "Responds to state changes without manual page reloads."),
            ("async fetch", "Supabase CRUD, AI chat, PC recommendations, exchange rates.", "Retrieves and updates external data."),
        ],
        widths=[1.3, 3.0, 2.1],
    )

    add_heading(doc, "Appendix B: Usability and Accessibility Test Checklist", 1)
    add_table(
        doc,
        ["Task / Check", "Result", "Notes"],
        [
            ("Open homepage on desktop, tablet, and mobile.", "Pass", "Hero content, navigation, and statistics adapt to screen size."),
            ("Search for a product from navigation.", "Pass", "Search field remains visible and usable across devices."),
            ("Add item to cart and view cart.", "Pass", "Cart state updates and data can be persisted through Supabase cart_items."),
            ("Add item to wishlist and view wishlist.", "Pass", "Wishlist state updates and data can be persisted through Supabase wishlist_items."),
            ("Complete checkout with stock validation.", "Pass", "System checks latest stock and deducts purchased quantity after order creation."),
            ("Use PC Builder after selecting first component.", "Partial", "AI recommendations depend on configured Groq key and quota; manual catalogue selection remains available."),
            ("Use AI chat.", "Partial", "Depends on configured Google API key and quota."),
            ("Keyboard through main navigation and forms.", "Partial", "Standard controls work; custom dropdowns/modals require deeper testing."),
            ("Check light and dark mode readability.", "Partial", "Most colours update dynamically; full contrast audit is recommended."),
            ("Review mobile admin pages.", "Pass", "Navigation compacts and cards/tables adapt, with horizontal table scrolling where needed."),
        ],
        widths=[2.6, 0.9, 2.9],
    )


def build():
    generate_wireframes()
    doc = Document()
    configure_document(doc)
    add_footer(doc)
    add_cover(doc)
    add_contents(doc)
    add_intro(doc)
    add_system_overview(doc)
    add_wireframes(doc)
    add_screenshots(doc)
    add_responsive(doc)
    add_usability(doc)
    add_accessibility(doc)
    add_improvements(doc)
    add_advanced(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendices(doc)
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
